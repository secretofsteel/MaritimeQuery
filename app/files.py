"""File helpers: hashing, JSONL cache, document reading and cleaning."""

from __future__ import annotations

import io
import json
import re
import pickle
import subprocess
import tempfile
from hashlib import md5
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple
import pandas as pd
from markdownify import markdownify as md
import yaml
import mammoth
import pymupdf.layout
import pymupdf4llm
import fitz  # PyMuPDF

from .logger import LOGGER
from .config import AppConfig

# ==============================================================================
#  HELPER FUNCTIONS
# ==============================================================================

def file_fingerprint(path: Path) -> Dict[str, Any]:
    """Return a simple fingerprint for change detection."""
    stat = path.stat()
    return {"name": path.name, "size": stat.st_size, "mtime": stat.st_mtime}

def iter_library_files(root: Path) -> Iterable[Path]:
    """Yield supported document paths from the library root."""
    for ext in ("*.docx", "*.doc", "*.xlsx", "*.xls", "*.pdf", "*.txt"):
        yield from root.glob(ext)

def current_files_index(root: Path) -> Dict[str, Dict[str, Any]]:
    """Build an index of current files keyed by filename."""
    return {p.name: file_fingerprint(p) for p in iter_library_files(root)}

def load_jsonl(path: Path) -> Dict[str, Any]:
    """Read a JSONL cache file into a dict keyed by filename."""
    data: Dict[str, Any] = {}
    if path.exists():
        for line in path.read_text(encoding="utf-8", errors="ignore").splitlines():
            try:
                record = json.loads(line)
                filename = record.get("filename")
                if filename:
                    data[filename] = record
            except json.JSONDecodeError:
                continue
    return data

def write_jsonl(path: Path, records: Iterable[Dict[str, Any]]) -> None:
    """Overwrite a JSONL file with records."""
    with path.open("w", encoding="utf-8") as file:
        for record in records:
            file.write(json.dumps(record, ensure_ascii=False) + "\n")

def upsert_jsonl_record(path: Path, record: Dict[str, Any]) -> None:
    """Update or append a record in a JSONL cache."""
    data = load_jsonl(path)
    data[record["filename"]] = record
    write_jsonl(path, data.values())

def _get_text_cache_path(file_path: Path, cache_dir: Path) -> Path:
    """Generate cache path using filename + hash."""
    file_stat = file_path.stat()
    cache_key = f"{file_stat.st_mtime}_{file_stat.st_size}"
    cache_hash = md5(cache_key.encode()).hexdigest()[:8]
    safe_name = re.sub(r'[^\w\-]', '_', file_path.stem)
    return cache_dir / f"{safe_name}_{cache_hash}.pkl"

def clean_text_for_llm(text: str) -> str:
    """Normalize text for LLMs without breaking basic Markdown structures."""
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    lines = text.split("\n")
    cleaned_lines = []
    in_code_block = False

    for line in lines:
        stripped = line.strip()

        # Track fenced code blocks ``` ```
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            cleaned_lines.append(line.rstrip())
            continue

        if in_code_block:
            # Do not touch code fences content
            cleaned_lines.append(line.rstrip())
            continue

        # For markdown table rows, keep pipes & hyphens as-is, just trim edges
        if stripped.startswith("|") and stripped.endswith("|"):
            # Collapse internal tabs/spaces a bit but keep structure
            norm = re.sub(r"[ \t]+", " ", line)
            cleaned_lines.append(norm.rstrip())
            continue

        # Generic normalization for normal text
        line = re.sub(r"[ \t]+", " ", line)          # collapse spaces
        line = re.sub(r"([\-_=]{4,})", "----", line) # normalize long ruler-ish runs
        line = re.sub(r"(\.{4,})", "...", line)      # normalize dot spam
        line = re.sub(r"([,.;:!?])\1{2,}", r"\1", line)  # kill stupid !!!??? spam

        cleaned_lines.append(line.rstrip())

    text = "\n".join(cleaned_lines)

    # Collapse excessive blank lines
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)

    return text.strip()

def is_scanned_pdf(path: Path) -> bool:
    """Detect if PDF is scanned (needs OCR) vs born-digital."""
    try:
        import fitz
    except ImportError:
        return False
    
    try:
        with fitz.open(str(path)) as doc:
            pages_to_check = min(3, len(doc))
            text_chars_total = 0
            
            for page_num in range(pages_to_check):
                page = doc[page_num]
                text = page.get_text("text") or ""
                text_chars_total += len(text.strip())
            
            return text_chars_total < 100
    except Exception as exc:
        LOGGER.warning("Failed to detect if %s is scanned: %s", path.name, exc)
        return False


# ==============================================================================
#  LIBREOFFICE CONVERSION
# ==============================================================================

def convert_to_pdf_with_libreoffice(input_path: Path, output_dir: Optional[Path] = None) -> Optional[Path]:
    """
    Convert DOCX/DOC to PDF using LibreOffice headless.
    Works on both Linux and Windows.
    
    Args:
        input_path: Path to input DOCX/DOC file
        output_dir: Optional output directory (uses temp if None)
        
    Returns:
        Path to converted PDF, or None if conversion failed
    """
    import sys
    
    if output_dir is None:
        output_dir = Path(tempfile.mkdtemp())
    else:
        output_dir.mkdir(parents=True, exist_ok=True)
    
    # Determine LibreOffice command based on platform
    if sys.platform == "win32":
        # Windows: Try common installation paths
        possible_paths = [
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            "soffice.exe",  # If in PATH
        ]
        libreoffice_cmd = None
        for path in possible_paths:
            if Path(path).exists() or path == "soffice.exe":
                libreoffice_cmd = path
                break
        
        if libreoffice_cmd is None:
            LOGGER.error("LibreOffice not found on Windows. Install from: https://www.libreoffice.org/download/")
            return None
    else:
        # Linux/Mac
        libreoffice_cmd = "libreoffice"
    
    try:
        # LibreOffice headless conversion
        result = subprocess.run(
            [
                libreoffice_cmd,
                "--headless",
                "--convert-to", "pdf",
                "--outdir", str(output_dir),
                str(input_path)
            ],
            capture_output=True,
            text=True,
            timeout=60  # 60 second timeout for large files
        )
        
        if result.returncode != 0:
            LOGGER.error("LibreOffice conversion failed for %s: %s", input_path.name, result.stderr)
            return None
        
        # Find the converted PDF
        pdf_name = input_path.stem + ".pdf"
        pdf_path = output_dir / pdf_name
        
        if pdf_path.exists():
            LOGGER.info("Successfully converted %s to PDF via LibreOffice", input_path.name)
            return pdf_path
        else:
            LOGGER.error("LibreOffice conversion succeeded but PDF not found: %s", pdf_path)
            return None
            
    except subprocess.TimeoutExpired:
        LOGGER.error("LibreOffice conversion timeout for %s (>60s)", input_path.name)
        return None
    except FileNotFoundError:
        if sys.platform == "win32":
            LOGGER.error("LibreOffice not found. Download from: https://www.libreoffice.org/download/")
        else:
            LOGGER.error("LibreOffice not found - install with: apt-get install libreoffice")
        return None
    except Exception as exc:
        LOGGER.error("LibreOffice conversion failed for %s: %s", input_path.name, exc)
        return None


# ==============================================================================
#  MAIN EXTRACTION LOGIC
# ==============================================================================

def read_doc_for_llm(path: Path, max_chars: Optional[int] = None, use_cache: bool = True) -> str:
    """
    Extract text from various document formats using Markdown-first approach.
    Uses modern libraries (Mammoth, PyMuPDF4LLM) with intelligent fallbacks.
    """
    if use_cache:
        config = AppConfig.get()
        cache_dir = config.paths.cache_dir / "text_extracts"
        cache_dir.mkdir(exist_ok=True)
        cache_path = _get_text_cache_path(path, cache_dir)
        
        if cache_path.exists():
            try:
                with open(cache_path, 'rb') as f:
                    cached_text = pickle.load(f)
                if max_chars and len(cached_text) > max_chars:
                    return cached_text[:max_chars]
                return cached_text
            except Exception as exc:
                LOGGER.warning("Failed to load text cache for %s: %s", path.name, exc)

    ext = path.suffix.lower()
    handlers = {
        '.txt': _extract_txt,
        '.docx': _extract_docx,
        '.doc': _extract_legacy_doc,
        '.pdf': _extract_pdf,
        '.xlsx': _extract_excel,
        '.xls': _extract_excel,
    }
    
    handler = handlers.get(ext)
    try:
        if handler:
            text = handler(path)
        else:
            LOGGER.warning("Unsupported file type: %s", ext)
            text = ""
    except Exception as exc:
        LOGGER.error("Failed to extract %s: %s", path.name, exc)
        text = f"[READ_ERROR] {exc}"

    # Cache result
    if use_cache and text and not text.startswith("[READ_ERROR]"):
        try:
            with open(cache_path, 'wb') as f:
                pickle.dump(text, f)
        except Exception:
            pass

    if max_chars and len(text) > max_chars:
        text = text[:max_chars]
    return text

def _extract_txt(path: Path) -> str:
    """Plain text extraction."""
    return path.read_text(encoding="utf-8", errors="ignore")

# ==============================================================================
#  DOCX EXTRACTION (Mammoth + Visual Extraction)
# ==============================================================================

def estimate_docx_pages(path: Path) -> int:
    """
    Estimate page count of a DOCX file.
    Uses heuristics: paragraphs + tables.
    
    Returns:
        Estimated page count (minimum 1)
    """
    try:
        doc = Document(str(path))
        
        # Count content elements
        paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
        table_count = len(doc.tables)
        
        # Heuristic: ~40 paragraphs per page, each table = 0.5 page
        estimated_pages = max(1, int(paragraph_count / 40 + table_count * 0.5))
        
        return estimated_pages
        
    except Exception as exc:
        LOGGER.warning("Failed to estimate pages for %s: %s", path.name, exc)
        return 10  # Default to "large" if we can't estimate


def _extract_docx(path: Path) -> str:
    """
    Smart DOCX extraction with intelligent routing:
    - Small files (<5 pages): Mammoth → Markdown (better for forms/checklists)
    - Large files (≥5 pages): LibreOffice → PDF → PyMuPDF4LLM (preserves numbers)
    - Fallback: python-docx if Mammoth fails
    """
    config = AppConfig.get()
    
    # Estimate page count
    page_count = estimate_docx_pages(path)
    LOGGER.info("Estimated %d pages for %s", page_count, path.name)
    
    # Route based on size
    if page_count < 5:
        # Small file - use Mammoth (better for forms/checklists)
        LOGGER.info("Using Mammoth for small file: %s", path.name)
        return _extract_docx_with_mammoth(path)
    else:
        # Large file - use LibreOffice conversion
        LOGGER.info("Using LibreOffice conversion for large file: %s", path.name)
        return _extract_docx_via_libreoffice(path)

def _extract_docx_with_mammoth(path: Path) -> str:
    """
    Extract small DOCX files with Mammoth (good for forms/checklists).
    """
    markdown_text = ""
    
    try:
        with path.open("rb") as docx_file:
            # Skip images - we handle them separately
            def ignore_image(image):
                return {}
            
            result = mammoth.convert_to_html(
                docx_file,
                convert_image=mammoth.images.img_element(ignore_image)
            )
            html_text = result.value
            
            if result.messages:
                for msg in result.messages:
                    if hasattr(msg, 'type') and msg.type == 'warning':
                        LOGGER.debug("Mammoth warning for %s: %s", path.name, msg.message)
            
            # Convert HTML to Markdown
            markdown_text = md(html_text, heading_style="ATX")
            
            if not markdown_text or len(markdown_text.strip()) < 50:
                LOGGER.warning("Mammoth output too short for %s, trying fallback", path.name)
                raise ValueError("Mammoth output too short")
                
    except Exception as exc:
        LOGGER.warning("Mammoth extraction failed for %s: %s, trying fallback", path.name, exc)
        try:
            markdown_text = _extract_docx_fallback(path)
        except Exception as fallback_exc:
            LOGGER.error("DOCX fallback also failed for %s: %s", path.name, fallback_exc)
            return f"[DOCX_EXTRACTION_FAILED] {exc}"
    
    # Visual extraction
    visuals_text = ""
    config = AppConfig.get()
    
    if getattr(config, 'visual_extraction_enabled', True):
        try:
            from docx import Document
            doc = Document(str(path))
            images = _extract_images_from_docx(doc)
            
            if images:
                visual_blocks = []
                for img in images:
                    block = f"\n\n[IMAGE - {img['context']}]\n{img['description']}\n[/IMAGE]\n"
                    visual_blocks.append(block)
                visuals_text = "\n".join(visual_blocks)
                LOGGER.info("Extracted %d visual descriptions from %s", len(images), path.name)
        except Exception as exc:
            LOGGER.warning("Visual extraction failed for %s: %s", path.name, exc)
    
    return markdown_text + visuals_text


def _extract_docx_via_libreoffice(path: Path) -> str:
    """
    Extract large DOCX files via LibreOffice → PDF → PyMuPDF4LLM.
    This preserves section numbering and complex formatting.
    """
    temp_dir = Path(tempfile.mkdtemp())
    
    try:
        # Convert to PDF
        pdf_path = convert_to_pdf_with_libreoffice(path, temp_dir)
        
        if pdf_path is None or not pdf_path.exists():
            LOGGER.error("LibreOffice conversion failed for %s, using fallback", path.name)
            return _extract_docx_fallback(path)
        
        # Extract from PDF
        markdown_text = _extract_pdf(pdf_path)
        
        # Clean up temp PDF
        try:
            pdf_path.unlink()
        except Exception:
            pass
        
        return markdown_text
        
    except Exception as exc:
        LOGGER.error("LibreOffice-based extraction failed for %s: %s", path.name, exc)
        return _extract_docx_fallback(path)
    finally:
        # Clean up temp directory
        try:
            temp_dir.rmdir()
        except Exception:
            pass


def _extract_docx_fallback(path: Path) -> str:
    """
    Fallback DOCX extraction using python-docx.
    Handles paragraphs, tables, and basic structure.
    """
    from docx import Document
    
    doc = Document(str(path))
    parts = []
    
    # Extract paragraphs with heading detection
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
            
        style_name = para.style.name if para.style else ""
        
        if style_name.startswith('Heading'):
            try:
                level = int(style_name.replace('Heading', '').strip())
                prefix = '#' * min(level, 6)
                parts.append(f"\n{prefix} {text}\n")
            except ValueError:
                parts.append(f"\n## {text}\n")
        elif style_name == 'Title':
            parts.append(f"\n# {text}\n")
        else:
            parts.append(text)
    
    # Extract tables
    for table in doc.tables:
        table_md = _convert_docx_table_to_markdown(table)
        if table_md:
            parts.append(f"\n\n{table_md}\n")
    
    return "\n\n".join(parts)


def _convert_docx_table_to_markdown(table) -> str:
    """Convert a python-docx table to Markdown format."""
    if not table.rows:
        return ""
    
    try:
        rows_text = []
        
        for row_idx, row in enumerate(table.rows):
            cells_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ').replace('|', '\\|')
                cells_text.append(cell_text)
            
            if cells_text:
                rows_text.append("| " + " | ".join(cells_text) + " |")
            
            # Add separator after first row (header)
            if row_idx == 0 and len(cells_text) > 0:
                separator = "|" + "|".join([" --- "] * len(cells_text)) + "|"
                rows_text.append(separator)
        
        return "\n".join(rows_text)
        
    except Exception as exc:
        LOGGER.debug("Failed to convert table to markdown: %s", exc)
        return ""

def _extract_images_from_docx(document) -> List[Dict[str, Any]]:
    """
    Extract embedded images from DOCX with context and descriptions.
    Filters out decorative elements (icons, logos) based on size.
    """
    from docx.oxml.ns import qn
    from PIL import Image
    
    images = []
    
    # Size thresholds to filter decorative elements
    MIN_WIDTH_PX = 100
    MIN_HEIGHT_PX = 100
    MIN_AREA_PX = 15000  # ~122x122 or larger
    
    for rel_id, rel in document.part.rels.items():
        if "image" not in rel.target_ref:
            continue
            
        try:
            image_part = rel.target_part
            image_bytes = image_part.blob
            
            # Check image dimensions to filter decorative elements
            with Image.open(io.BytesIO(image_bytes)) as pil_img:
                width, height = pil_img.size
                area = width * height
                
                # Skip small images (logos, icons, decorative elements)
                if width < MIN_WIDTH_PX or height < MIN_HEIGHT_PX or area < MIN_AREA_PX:
                    LOGGER.debug("Skipping small image: %dx%d (area=%d)", width, height, area)
                    continue
            
            # Extract context from surrounding text
            context = _extract_image_context_docx(document, rel_id)
            
            # Get AI description of the image
            description = _describe_visual_with_gemini(image_bytes, context)
            
            if description:
                images.append({
                    'type': 'IMAGE',
                    'context': context or f"Image {len(images)+1}",
                    'description': description,
                    'size': f"{width}x{height}"
                })
                
        except Exception as exc:
            LOGGER.debug("Failed to process image in DOCX: %s", exc)
            continue
    
    return images

def _extract_image_context_docx(document, rel_id: str) -> str:
    """Extract text context around an image in DOCX."""
    from docx.oxml.ns import qn
    
    # Find paragraphs containing this image
    context_parts = []
    
    for i, paragraph in enumerate(document.paragraphs):
        # Check if paragraph contains the image
        for run in paragraph.runs:
            if run._element.xpath(f'.//a:blip[@r:embed="{rel_id}"]'):
                # Found the paragraph with the image
                # Collect surrounding context
                start_idx = max(0, i - 2)
                end_idx = min(len(document.paragraphs), i + 3)
                
                for ctx_para in document.paragraphs[start_idx:end_idx]:
                    text = ctx_para.text.strip()
                    if text and len(text) > 10:  # Skip empty or very short lines
                        context_parts.append(text)
                
                break
    
    # Also check for captions or titles in nearby text
    context = " | ".join(context_parts) 
    return context if context else ""  

# ==============================================================================
#  PDF EXTRACTION (PyMuPDF4LLM + Fallback + Visual Extraction)
# ==============================================================================

def _extract_pdf(path: Path) -> str:
    """
    Extract from PDF using PyMuPDF4LLM (layout-aware Markdown conversion).
    Falls back to custom extraction if PyMuPDF4LLM produces poor output.
    Includes optional visual extraction of diagrams and images.
    
    Now includes:
    - Better table detection (table_strategy="lines")
    - Progress indicator for user feedback
    - Header/footer exclusion for cleaner extraction
    """
    # 1. Primary: PyMuPDF4LLM for layout-aware Markdown
    markdown_text = ""
    extraction_method = "pymupdf4llm"
    
    try:        
        doc = pymupdf.open(str(path))
        
        markdown_text = pymupdf4llm.to_markdown(
            doc,
            header=False,           # Exclude headers (intelligent detection)
            footer=False,           # Exclude footers (intelligent detection)
            table_strategy="lines", # More aggressive table detection
            show_progress=True      # Show progress bar for user feedback
        )
        
        # Quality check: if output is suspiciously short or garbled, trigger fallback
        expected_min_length = path.stat().st_size // 100  # Rough heuristic
        if len(markdown_text) < expected_min_length or len(markdown_text) < 500:
            LOGGER.warning("PyMuPDF4LLM output suspiciously short for %s (%d chars), trying fallback", 
                          path.name, len(markdown_text))
            raise ValueError("Output too short")
            
    except Exception as exc:
        LOGGER.warning("PyMuPDF4LLM failed for %s: %s, using custom fallback", path.name, exc)
        extraction_method = "custom_fallback"
        
        # Fallback: Custom extraction with layout awareness
        try:
            markdown_text = _extract_pdf_with_layout_fallback(path)
        except Exception as fallback_exc:
            LOGGER.error("PDF fallback extraction also failed for %s: %s", path.name, fallback_exc)
            # Last resort: plain text extraction
            try:
                with fitz.open(str(path)) as doc:
                    markdown_text = "\n\n".join([page.get_text() for page in doc])
                extraction_method = "plain_text"
            except Exception:
                return f"[PDF_EXTRACTION_FAILED] {exc}"

    # 2. Visual extraction (Diagrams/Images) - Optional
    config = AppConfig.get()
    visuals_text = ""
    
    if getattr(config, 'visual_extraction_enabled', True):
        try:
            visual_items = []
            with fitz.open(str(path)) as doc:
                for page_num, page in enumerate(doc):
                    items = _extract_visual_content_from_page(page, page_num)
                    visual_items.extend(items)
            
            if visual_items:
                visual_blocks = []
                for item in visual_items:
                    v_type = item.get('type', 'VISUAL').upper()
                    block = f"\n\n> **[{v_type}: {item['context']}]**\n> {item['description']}\n"
                    visual_blocks.append(block)
                
                visuals_text = "\n".join(visual_blocks)
                LOGGER.info("Extracted %d visual descriptions from %s", len(visual_items), path.name)
                
        except Exception as exc:
            LOGGER.warning("Visual extraction failed for %s: %s", path.name, exc)

    LOGGER.info("PDF %s extracted via %s", path.name, extraction_method)
    return markdown_text + visuals_text

def _extract_pdf_with_layout_fallback(path: Path) -> str:
    """
    Custom PDF extraction with layout awareness for multi-column documents.
    Used as fallback when PyMuPDF4LLM fails or produces poor output.
    """
    with fitz.open(str(path)) as doc:
        all_text = []
        
        for page_num, page in enumerate(doc):
            # Use dict mode for better structure
            page_dict = page.get_text("dict")
            blocks = page_dict.get("blocks", [])
            
            # Sort blocks by position (Y then X) to handle columns
            text_blocks = [b for b in blocks if b.get("type") == 0]  # Type 0 = text
            
            # Detect columns by clustering X positions
            if text_blocks:
                sorted_blocks = _sort_blocks_with_column_detection(text_blocks, page.rect.width)
                
                page_text = []
                for block in sorted_blocks:
                    for line in block.get("lines", []):
                        line_text = ""
                        for span in line.get("spans", []):
                            line_text += span.get("text", "")
                        if line_text.strip():
                            page_text.append(line_text)
                
                all_text.append("\n".join(page_text))
        
        return "\n\n".join(all_text)

def _sort_blocks_with_column_detection(blocks: List[Dict], page_width: float) -> List[Dict]:
    """
    Sort text blocks considering multi-column layouts.
    Detects column breaks and reads top-to-bottom within each column.
    """
    if not blocks:
        return blocks
    
    # Extract X positions of all blocks
    x_positions = [b["bbox"][0] for b in blocks]
    
    # Simple column detection: if there's a significant gap in X positions
    x_positions_sorted = sorted(set(x_positions))
    
    # Detect column break (gap > 20% of page width)
    column_threshold = page_width * 0.2
    columns = []
    current_column = [x_positions_sorted[0]]
    
    for x in x_positions_sorted[1:]:
        if x - current_column[-1] > column_threshold:
            columns.append(current_column)
            current_column = [x]
        else:
            current_column.append(x)
    columns.append(current_column)
    
    # If multi-column detected, sort within columns
    if len(columns) > 1:
        sorted_blocks = []
        for col_x_range in columns:
            col_min = min(col_x_range)
            col_max = max(col_x_range) + column_threshold
            
            # Get blocks in this column
            col_blocks = [b for b in blocks if col_min <= b["bbox"][0] < col_max]
            # Sort by Y position within column
            col_blocks.sort(key=lambda b: b["bbox"][1])
            sorted_blocks.extend(col_blocks)
        
        return sorted_blocks
    else:
        # Single column: simple Y-then-X sort
        blocks.sort(key=lambda b: (b["bbox"][1], b["bbox"][0]))
        return blocks

def _extract_image_context_pdf_by_bbox(page, page_num: int, img_bbox: fitz.Rect) -> str:
    """Extract text context around an image using its bounding box."""
    blocks = page.get_text("dict")["blocks"]
    text_blocks = [b for b in blocks if b.get("type") == 0]
    
    context_parts = []
    for block in text_blocks:
        block_rect = fitz.Rect(block["bbox"])
        
        # Check if block is near the image (within 50 points)
        if block_rect.y1 < img_bbox.y0 - 50:  # Too far above
            continue
        if block_rect.y0 > img_bbox.y1 + 50:  # Too far below
            continue
        
        # Extract text from block
        block_text = ""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                block_text += span.get("text", "")
        
        if block_text.strip() and len(block_text.strip()) > 10:
            context_parts.append(block_text.strip())
    
    context = " | ".join(context_parts[:2])
    return context if context else f"Page {page_num+1}"

def _extract_visual_content_from_page(page, page_num: int) -> List[Dict[str, Any]]:
    """Extract significant visual elements from PDF page."""
    visual_items = []
    
    # Size thresholds (in page points, not pixels)
    MIN_WIDTH_PT = 100
    MIN_HEIGHT_PT = 100
    MIN_AREA_PT = 15000
    
    # 1. Extract images with bbox and xref in one call
    image_infos = page.get_image_info(xrefs=True)
    
    for img_index, info in enumerate(image_infos):
        xref = info['xref']
        bbox = fitz.Rect(info['bbox'])
        
        # Filter by RENDERED size on page (catches scaled-down high-res images)
        if bbox.width < MIN_WIDTH_PT or bbox.height < MIN_HEIGHT_PT:
            continue
        
        if bbox.width * bbox.height < MIN_AREA_PT:
            continue
        
        try:
            base_image = page.parent.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Validate actual pixel dimensions too (double-check)
            from PIL import Image
            with Image.open(io.BytesIO(image_bytes)) as pil_img:
                width, height = pil_img.size
                # If pixel size is tiny, skip (even if rendered large)
                if width < 100 or height < 100:
                    continue
            
            # Extract context using the bbox we already have
            context = _extract_image_context_pdf_by_bbox(page, page_num, bbox)
            
            description = _describe_visual_with_gemini(image_bytes, context)
            
            if description:
                visual_items.append({
                    'type': 'IMAGE',
                    'context': context or f"Page {page_num+1}, Image {img_index+1}",
                    'description': description,
                    'size': f"{width}x{height}px, rendered {int(bbox.width)}x{int(bbox.height)}pt"
                })
                
        except Exception as exc:
            LOGGER.debug("Failed to extract image from page %d: %s", page_num, exc)
            continue
    
    # 2. Extract vector drawings (diagrams, charts)
    try:
        drawings = page.cluster_drawings()
        for drawing_idx, drawing_cluster in enumerate(drawings):
            if not drawing_cluster:
                continue
            
            # Calculate bounding box of the drawing cluster
            bbox = fitz.Rect()
            for item in drawing_cluster:
                if hasattr(item, 'rect'):
                    bbox |= item.rect
            
            width = bbox.width
            height = bbox.height
            area = width * height
            
            # Filter small drawings (decorative lines, bullets, etc.)
            if width < MIN_WIDTH or height < MIN_HEIGHT or area < MIN_AREA:
                continue
            
            # Render the drawing area as image
            try:
                mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for better quality
                pix = page.get_pixmap(matrix=mat, clip=bbox)
                image_bytes = pix.tobytes("png")
                
                # Extract context
                context = _extract_drawing_context_pdf(page, page_num, bbox)
                
                # Get AI description
                description = _describe_visual_with_gemini(image_bytes, context)
                
                if description:
                    visual_items.append({
                        'type': 'DIAGRAM',
                        'context': context or f"Page {page_num+1}, Diagram {drawing_idx+1}",
                        'description': description,
                        'size': f"{int(width)}x{int(height)}"
                    })
                    
            except Exception as exc:
                LOGGER.debug("Failed to render drawing from page %d: %s", page_num, exc)
                continue
                
    except Exception as exc:
        LOGGER.debug("Failed to cluster drawings on page %d: %s", page_num, exc)
    
    return visual_items

def _extract_image_context_pdf(page, page_num: int, img_index: int) -> str:
    """Extract text context around an image in PDF."""
    # Get all text blocks on the page
    blocks = page.get_text("dict")["blocks"]
    text_blocks = [b for b in blocks if b.get("type") == 0]
    
    # Get image rectangles
    img_list = page.get_image_info()
    if img_index >= len(img_list):
        return f"Page {page_num+1}"
    
    img_rect = fitz.Rect(img_list[img_index]["bbox"])
    
    # Find text blocks near the image (within 50 points)
    context_parts = []
    for block in text_blocks:
        block_rect = fitz.Rect(block["bbox"])
        # Check if block is above, below, or beside the image
        if block_rect.y1 < img_rect.y0 - 50:  # Above
            continue
        if block_rect.y0 > img_rect.y1 + 50:  # Below
            continue
        
        # Extract text from block
        block_text = ""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                block_text += span.get("text", "")
        
        if block_text.strip() and len(block_text.strip()) > 10:
            context_parts.append(block_text.strip())
    
    context = " | ".join(context_parts)
    return context if context else f"Page {page_num+1}"

def _extract_drawing_context_pdf(page, page_num: int, bbox: fitz.Rect) -> str:
    """Extract text context around a drawing in PDF."""
    blocks = page.get_text("dict")["blocks"]
    text_blocks = [b for b in blocks if b.get("type") == 0]
    
    context_parts = []
    for block in text_blocks:
        block_rect = fitz.Rect(block["bbox"])
        
        # Check proximity to drawing
        if block_rect.y1 < bbox.y0 - 50 or block_rect.y0 > bbox.y1 + 50:
            continue
        
        # Extract text
        block_text = ""
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                block_text += span.get("text", "")
        
        if block_text.strip() and len(block_text.strip()) > 10:
            context_parts.append(block_text.strip())
    
    context = " | ".join(context_parts)
    return context if context else f"Page {page_num+1}"

# ==============================================================================
#  EXCEL EXTRACTION (Smart YAML/Markdown switching)
# ==============================================================================

def _extract_excel(path: Path) -> str:
    """
    Extract from Excel with smart format selection:
    - Wide tables (>5 columns) → YAML for better LLM comprehension
    - Narrow tables (≤5 columns) → Markdown for readability
    """
    try:
        excel = pd.ExcelFile(str(path))
    except Exception as exc:
        LOGGER.error("Failed to open Excel file %s: %s", path.name, exc)
        return f"[EXCEL_READ_ERROR] {exc}"
    
    chunks: List[str] = []
    
    for sheet_name in excel.sheet_names[:5]:  # Limit to first 5 sheets
        try:
            df = pd.read_excel(str(path), sheet_name=sheet_name, nrows=100, header=None)
            
            # Clean up empty rows/columns
            df.dropna(axis=1, how="all", inplace=True)
            df.dropna(thresh=2, inplace=True)
            
            if df.empty:
                continue
            
            # Heuristic: Detect headers (first row likely contains column names)
            try:
                df.columns = df.iloc[0]
                df = df[1:].reset_index(drop=True)
            except Exception:
                pass
            
            # Strategy selection based on table width
            num_columns = len(df.columns)
            
            if num_columns > 5:
                # WIDE TABLE: Convert to YAML for better key-value preservation
                records = df.to_dict(orient='records')
                
                # Clean nulls for readable YAML
                clean_records = []
                for record in records[:50]:  # Limit rows to prevent token bloat
                    clean_record = {k: v for k, v in record.items() if pd.notna(v)}
                    if clean_record:  # Only include non-empty records
                        clean_records.append(clean_record)
                
                if clean_records:
                    yaml_text = yaml.dump(clean_records, sort_keys=False, 
                                         default_flow_style=False, allow_unicode=True)
                    chunks.append(f"## Sheet: {sheet_name}\n\n```yaml\n{yaml_text}\n```")
                    LOGGER.debug("Sheet '%s' in %s: wide table (%d cols) → YAML", 
                               sheet_name, path.name, num_columns)
            else:
                # NARROW TABLE: Markdown is fine and more readable
                try:
                    md = df.to_markdown(index=False)
                    chunks.append(f"## Sheet: {sheet_name}\n\n{md}")
                    LOGGER.debug("Sheet '%s' in %s: narrow table (%d cols) → Markdown", 
                               sheet_name, path.name, num_columns)
                except Exception:
                    # Fallback if markdown conversion fails
                    csv_text = df.to_csv(index=False)
                    chunks.append(f"## Sheet: {sheet_name}\n\n```csv\n{csv_text}\n```")
        
        except Exception as exc:
            LOGGER.warning("Failed to extract sheet '%s' from %s: %s", 
                          sheet_name, path.name, exc)
            continue
    
    if not chunks:
        return "[EXCEL_NO_CONTENT]"
    
    return "\n\n".join(chunks)

# ==============================================================================
#  LEGACY DOC SUPPORT
# ==============================================================================

def _extract_legacy_doc(path: Path) -> str:
    """
    Extract from legacy .doc files via LibreOffice → PDF → PyMuPDF4LLM.
    """
    LOGGER.info("Processing legacy .doc file via LibreOffice: %s", path.name)
    return _extract_docx_via_libreoffice(path)

# ==============================================================================
#  VISUAL EXTRACTION HELPERS
# ==============================================================================

def _describe_visual_with_gemini(image_bytes: bytes, context: str = "") -> Optional[str]:
    """
    Send visual element to Gemini for description.
    Uses Flash Lite for cost-effectiveness on diagrams/images.
    """
    try:
        from google.genai import types
        from PIL import Image
        
        config = AppConfig.get()
        
        # Verify image is valid before sending
        try:
            with Image.open(io.BytesIO(image_bytes)) as img:
                # Convert to RGB if needed (some PDFs have CMYK)
                if img.mode not in ('RGB', 'L'):
                    img = img.convert('RGB')
                
                # Re-encode to ensure clean PNG
                output = io.BytesIO()
                img.save(output, format='PNG')
                image_bytes = output.getvalue()
        except Exception as img_exc:
            LOGGER.debug("Failed to validate/convert image: %s", img_exc)
            return None
        
        # Build prompt
        prompt = (
            "Describe this technical diagram/chart/drawing/table/image in detail. "
            "Focus on: labels, measurements, flow/connections, components, visual relationships,"
            "units, and key operational information. "
            "Be specific and technical. Use bullet point and/or table structure if required."
            "Ensure to keep it to the point, including all relevant details. Avoid flowery language."
        )
        if context:
            prompt += f"\n\nContext from document: {context}"
        
        response = config.client.models.generate_content(
            model="gemini-2.5-flash-lite",
            contents=[
                types.Part(text=prompt),
                types.Part(
                    inline_data=types.Blob(
                        mime_type="image/png",
                        data=image_bytes
                    )
                )
            ],
            config=types.GenerateContentConfig(
                temperature=0.2,  # Lower temp for technical descriptions
                #max_output_tokens=400  # Cap to prevent bloat
            )
        )
        
        description = response.text.strip() if response.text else None
        
        if description:
            LOGGER.debug("Generated visual description (%d chars)", len(description))
            return description
        
        return None
        
    except Exception as exc:
        LOGGER.warning("Gemini vision description failed: %s", exc)
        return None

__all__ = [
    "file_fingerprint",
    "iter_library_files",
    "current_files_index",
    "load_jsonl",
    "write_jsonl",
    "upsert_jsonl_record",
    "clean_text_for_llm",
    "read_doc_for_llm",
    "is_scanned_pdf",
]
